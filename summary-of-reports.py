import os
import re
from win32com import client as wc
import xlrd
from xlutils.copy import copy
from docx import Document

# 获取目标文件.doc列表
def file_name(name,pat,path_list):
    try:
        pires=os.listdir(name)
        for i in range(len(pires)):
            pires_name=name+'/'+pires[i]
            pire=os.listdir(pires_name)
            for j in pire:
                if pat.search(j):
                    path_list.append(pires_name+'/'+j)
        return path_list
    except:
        print('报告文件绝对路径出错')

#把.doc转化为.docx,保存到D:/Desktop/文档操作/
def doc2docx(doc_name, docx_name):
    if re.search(r'x$',doc_name):
        doc=Document(doc_name)
        doc.save(docx_name)
    else:
        try:
            word = wc.Dispatch("Word.Application")  # 打开word应用程序
            doc = word.Documents.Open(doc_name)  # 打开word文件
            doc.SaveAs(docx_name, 16)  # 另存为后缀为".docx"的文件，其中参数12指docx文件
            doc.Close()  # 关闭原来word文件
            word.Quit()
        except:
            print('转换文件错误，请检查文档默认打开软件是否为word')

#把列表中的.doc文件改为.docx,并返回.docx文件名列表
def change(lis,docxs_list,wook_file):
    for i in range(len(lis)):
        docn=lis[i]
        docxn=wook_file+'/'+str(i)+'.docx'
        docxs_list.append(docxn)
        doc2docx(docn,docxn)
    return docxs_list

#将.docx文件内容加入.xls
def set_excel(docxna,xlsna,zuzhang_name,ji_tuan,xu_hao,ptb_xuhao):
    doc = Document(docxna)
    tables = doc.tables
    get_number = re.compile(r'\d')
    b0 = xu_hao
    b1 = '康布尔'
    # 报告编号
    bianhao = tables[0].rows[0].cells[4]
    biao_hao = bianhao.text  # 报告编号：012021131
    a2 = re.search(r'\d+', biao_hao)
    b2 = a2.group()

    # 资质领导小组
    for table in tables:
        if table.rows[0].cells[0].text == '受评估队伍':
            lingdaoxiaozu = table.rows[2].cells[2]
            a3 = lingdaoxiaozu.text
            if a3.endswith('资质初审领导小组'):
                b3 = a3.rstrip('资质初审领导小组')
            elif a3.endswith('初审领导小组'):
                b3=a3.rstrip('初审领导小组')

            # 单位名称
            danwei = table.rows[1].cells[1]
            b4 = danwei.text

            # 队伍名称
            duiwu = table.rows[0].cells[1]
            b5 = duiwu.text

    # 专业类别
    zhuanye = tables[0].rows[6].cells[2]
    b6 = zhuanye.text

    # 评估类别
    pinggu = tables[0].rows[7].cells[2]
    b7 = pinggu.text

    # 资质等级
    for table in tables:
        if table.rows[0].cells[0].text == '受评估队伍':
            dengji = table.rows[1].cells[-1]
            b8 = dengji.text

    for table in tables:
        if len(table.rows[0].cells) > 8:
            # 设备配套表
            if table.rows[0].cells[3].text == '型 号':
                # 设备型号
                if b6.startswith('钻'):
                    for par in doc.paragraphs:
                        neirong = par.text
                        if re.search(r'ZJ\d{2}[A-Z]*', neirong):
                            shebeixinghao = re.search(r'ZJ\d{2}[A-Z]*', neirong)
                            b9 = shebeixinghao.group()
                            break
                        else:
                            continue
                else:
                    b9 = table.rows[1].cells[3].text

                # 设备生产厂家
                b10 = table.rows[1].cells[7].text

                # 设备出厂日期
                b11 = table.rows[1].cells[6].text

                # 设备投产日期
                b12 = b11

                # 井架型号
                for row in table.rows:
                    if row.cells[2].text == '井架':
                        b13 = row.cells[3].text

                # 井架形式
                if b6.startswith('钻'):
                    b14 = 'K型'
                elif re.search(r'XT',b13):
                    b14 = 'BJ'
                else:
                    b14 = '桅杆式'

                # 井架生产厂家
                for row in table.rows:
                    if row.cells[2].text == '井架':
                        b15 = row.cells[7].text

                        # 井架出厂日期
                        b16 = row.cells[6].text
                        # 投产日期
                        b17 = b16
            else:
                continue
        else:
            continue
    # 井架应力测试表格
    for table in tables:
        if len(table.rows[0].cells) > 2:
            if table.rows[0].cells[1].text == '井架型号规格':
                # 井架设计钩载
                gouzai = re.search(r'\d{2,3}', b13)
                gouzai1 = gouzai.group()
                gouzai2 = gouzai1 + '0'
                shijigozai = table.rows[8].cells[2].text
                lunxi = re.search(r'（.+）', shijigozai)
                lunxi1 = lunxi.group()
                if gouzai2 == '600':
                    b18 = '585kN' + lunxi1
                else:
                    b18 = gouzai2 + 'kN' + lunxi1

                # 井架承载能力
                chengzai = re.search(r'\d{3,4}kN', shijigozai)
                a19 = chengzai.group()
                b19 = a19 + lunxi1

                # 井架级别
                jiebie = table.rows[10].cells[2].text
                jiebie1 = re.search(r'[A-D]级', jiebie)
                b20 = jiebie1.group()

                # 井架得分
                defen = table.rows[11].cells[2].text
                a21 = re.search(r'\d{3}', defen)
                b21 = a21.group()

    # 得分表
    if len(tables[-1].rows) > 3:
        table_last = tables[-1]
    else:
        table_last = tables[-2]
    table_last_num = len(table_last.rows)
    # 无损检测得分
    li1_num = 0
    a22 = 0
    for cell in table_last.rows[0].cells:
        li1_num += 1
        if cell.text.startswith('无损'):
            for i in range(1, table_last_num - 1):
                jiancefen = table_last.rows[i].cells[li1_num - 1]
                if get_number.search(jiancefen.text):
                    aa22 = jiancefen.text.replace(' ', '')
                    a22 += float(aa22)
            b22 = str(a22)
    # 评估检查表得分
    li2_num = 0
    a23 = 0
    for cell in table_last.rows[0].cells:
        li2_num += 1
        if cell.text.startswith('评估'):
            for i in range(1, table_last_num - 1):
                jianchafen = table_last.rows[i].cells[li2_num - 1]
                if get_number.search(jianchafen.text):
                    a23 += float(jianchafen.text.replace(' ', ''))
        b23 = str(a23)

    # 技术水平得分
    a24 = 0
    if table_last.rows[0].cells[5].text.startswith('技术'):
        for i in range(1, len(table_last.rows) - 1):
            if get_number.search(table_last.rows[i].cells[5].text):
                a24 += float(table_last.rows[i].cells[5].text.replace(' ', ''))
        b24 = str(a24)
    else:
        b24 = '/'

    # 设备新旧得分
    a25 = 0
    for i in range(1, len(table_last.rows) - 1):
        if get_number.search(table_last.rows[i].cells[-3].text):
            a25 += float(table_last.rows[i].cells[-3].text.replace(' ', ''))
    b25 = str(a25)

    # 配套得分
    a26 = 0
    for i in range(1, len(table_last.rows) - 1):
        if get_number.search(table_last.rows[i].cells[-2].text):
            a26 += float(table_last.rows[i].cells[-2].text.replace(' ', ''))
    b26 = str(a26)

    # 综合得分
    b27 = table_last.rows[-1].cells[1].text

    # 评估级别
    a27 = float(b27)
    if a27 >= 900:
        b28 = 'Ⅰ'
    elif 900 > a27 >= 750:
        b28 = 'Ⅱ'
    elif 750 > a27 >= 600:
        b28 = 'Ⅲ'
    else:
        b28 = 'Ⅳ'

    # 评估结论
    if b28 == 'Ⅳ':
        b29 = '报废'
    else:
        b29 = '继续使用'

    # 评估时间
    for table in tables:
        if table.rows[0].cells[0].text == '受评估队伍':
            b30 = table.rows[6].cells[1].text

    # 评估组长
    b31 = zuzhang_name

    # 报告签发日期
    b32 = ''

    # 报告签发人
    b33 = '刘江涛'

    # 所在第区
    for table in tables:
        if table.rows[0].cells[0].text == '受评估队伍':
            b34 = table.rows[6].cells[-1].text

    # 集团内外
    b35 = ji_tuan

    # 备注
    b36 = ''
    ans_lis = [b0, b1, b2, b3, b4, b5, b6, b7, b8, b9, b10, b11, b12, b13, b14, b15, b16, b17, b18, b19, b20, b21, b22,
               b23, b24, b25, b26, b27, b28, b29, b30, b31, b32, b33, b34, b35, b36]
    # 写入Excel
    wb = xlrd.open_workbook(xlsna, formatting_info=True)
    xwb = copy(wb)
    if b35 == '集团内':
        sheet1 = xwb.get_sheet('计划内评估结果')
        sheet2 = xwb.get_sheet('主要设备信息')
    else:
        sheet1 = xwb.get_sheet('计划外评估结果')
        sheet2 = xwb.get_sheet('计划外设备信息')
    # 评估结果统计表
    for i in range(36):
        sheet1.write(xu_hao, i, ans_lis[i])

    # 设备配套统计
    for table in tables:
        if len(table.rows[0].cells) > 8:
            # 设备配套表
            if table.rows[0].cells[3].text == '型 号':
                ptb_row = len(table.rows) - 1
                for i in range(ptb_row):
                    sheet2.write(ptb_xuhao + i, 0, b2)
                    sheet2.write(ptb_xuhao + i, 1, b5)
                    for j in range(6):
                        sheet2.write(ptb_xuhao + i, j + 2, table.rows[i + 1].cells[2 + j].text)
    xwb.save(xlsna)
    return ptb_row
def main():
    pat = re.compile(r'\d{9,10}.+\.doc')
    ptb_xuhao=1
    path_list = []#目标文集路径列表
    docxs_list=[]
    print('==============-评估部-评估报告统计系统==============')
    print('温馨提示：')
    print('1、程序运行前请保证文档(.doc/.docx)默认已word打开,用word打开文档后出现弹窗勾选‘不在提示’')
    print('2、在桌面创建一个文件夹并把统计表放入文件夹中，统计表要删除内容，表头不删')
    print('3、请将需要统计的评估报告按计划内、外分开保存到两个文件夹，如已实现请忽略')
    print('4、如有疑问请联系：18829789062')
    zuzhang_name=input('组长姓名:')
    ji_tuan1=input('计划内报告请输入‘集团内’，计划外请输入‘集团外’:')
    ji_tuan=ji_tuan1.replace('\\','/')
    xu_hao=0
    jindu=0
    wook_file1=input('在桌面创建一个文件夹并把统计表放入文件夹中，输入文件夹的绝对路径:')
    wook_file=wook_file1.replace('\\','/')
    xlsna1 = input('请输入统计表名称带.xls:')
    xlsna2=xlsna1.replace('\\','/')
    xlsna=wook_file+'/'+ xlsna2
    # filena='D:/Desktop/工作集合/2021年设备评估/2021计划外评估'#文件夹路径
    filena1=input('请输入评估报告文件夹的绝对路径:')
    filena=filena1.replace('\\','/')
    file_name(filena,pat,path_list)
    change(path_list,docxs_list,wook_file)
    print('文件生成完成')
    for docxna in docxs_list:
        xu_hao += 1
        jindu += 1
        print('\r当前进度：{:.2f}%'.format(jindu*100/len(docxs_list)),end='')
        try:
            set_excel(docxna, xlsna, zuzhang_name, ji_tuan, xu_hao, ptb_xuhao)
            ptb_row = set_excel(docxna, xlsna, zuzhang_name, ji_tuan, xu_hao, ptb_xuhao)
            ptb_xuhao += ptb_row
        except:
            xu_hao = xu_hao-1
            continue
main()
